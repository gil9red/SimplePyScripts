#!/usr/bin/env python3
# -*- coding: utf-8 -*-

__author__ = "ipetrash"


# pip install python-docx
import docx


# SOURCE: https://github.com/python-openxml/python-docx/issues/74#issuecomment-261169410
def add_hyperlink(paragraph, url, text, color="1111FF", underline=True):
    """
    A function that places a hyperlink within a paragraph object.

    :param paragraph: The paragraph we are adding the hyperlink to.
    :param url: A string containing the required url
    :param text: The text displayed for the url
    :param color: color
    :param underline: underline
    :return: The hyperlink object
    """

    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(
        url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True
    )

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement("w:hyperlink")
    hyperlink.set(
        docx.oxml.shared.qn("r:id"),
        r_id,
    )

    # Create a w:r element
    new_run = docx.oxml.shared.OxmlElement("w:r")

    # Create a new w:rPr element
    rPr = docx.oxml.shared.OxmlElement("w:rPr")

    # Add color if it is given
    if not color is None:
        c = docx.oxml.shared.OxmlElement("w:color")
        c.set(docx.oxml.shared.qn("w:val"), color)
        rPr.append(c)

    # Remove underlining if it is requested
    if not underline:
        u = docx.oxml.shared.OxmlElement("w:u")
        u.set(docx.oxml.shared.qn("w:val"), "none")
        rPr.append(u)

    u = docx.oxml.shared.OxmlElement("w:u")
    u.set(docx.oxml.shared.qn("w:val"), "single")
    rPr.append(u)

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)

    return hyperlink


if __name__ == "__main__":
    document = docx.Document()
    document.add_heading("Hyperlink", level=0)

    p = document.add_paragraph("hyperlink: ")
    add_hyperlink(p, "https://google.ru", "google")

    p = document.add_paragraph("hyperlink: ")
    add_hyperlink(p, "https://google.ru", "google", underline=False)

    p = document.add_paragraph("hyperlink: ")
    add_hyperlink(p, "https://google.ru", "google", color="00FF00")

    p = document.add_paragraph("hyperlink: ")
    add_hyperlink(p, "https://google.ru", "google", color="FF8822")

    # Empty line
    document.add_paragraph()

    p = document.add_paragraph("Hello World! -> ")
    add_hyperlink(p, "https://google.ru/search?q=Hello", "Hello")
    p.add_run(" ")
    add_hyperlink(p, "https://google.ru/search?q=World", "World")
    p.add_run(" -> ")
    add_hyperlink(p, "https://google.ru/search?q=Hello World!", "Hello World!")

    # Save
    file_name_doc = "word.docx"
    document.save(file_name_doc)

    # Open file
    import os

    os.startfile(file_name_doc)
