#!/usr/bin/env python3
# -*- coding: utf-8 -*-

__author__ = "ipetrash"


from docx import Document


def create_table(document, headers, rows, style="Table Grid"):
    cols_number = len(headers)

    table = document.add_table(rows=1, cols=cols_number)
    table.style = style

    hdr_cells = table.rows[0].cells
    for i in range(cols_number):
        hdr_cells[i].text = headers[i]

    for row in rows:
        row_cells = table.add_row().cells
        for i in range(cols_number):
            row_cells[i].text = str(row[i])

    return table


document = Document()

headers = ("№ п/п", "Наименование параметра", "Единицы измерения", "Значение")
records_table1 = (
    (0, "Nan", "Nan", 0),
    (1, "Первая величина", "-/-", 0),
    (2, "Вторая величина", "-/-", "Базальт"),
    (3, "Третья величина", "м^2/ч", 0),
)
table1 = create_table(document, headers, records_table1)

document.add_paragraph()

rows = [[x, x, x * x] for x in range(1, 10)]
table2 = create_table(document, ("x", "y", "x * y"), rows)

document.save("testing.docx")
