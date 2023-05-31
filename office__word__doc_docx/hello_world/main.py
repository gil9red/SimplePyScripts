#!/usr/bin/env python3
# -*- coding: utf-8 -*-

__author__ = "ipetrash"


# SOURCE: https://github.com/python-openxml/python-docx
# SOURCE: https://python-docx.readthedocs.io/en/latest/


import os

# pip install python-docx
import docx
from docx.shared import Inches, Cm, Mm


document = docx.Document()
document.add_heading("Document Title", level=0)

p = document.add_paragraph("A plain paragraph having some ")
p.add_run("bold").bold = True
p.add_run(" and some ")
p.add_run("italic.").italic = True

document.add_heading("Heading, level 1", level=1)
document.add_paragraph("Intense quote", style="Intense Quote")
document.add_paragraph("first item in unordered list", style="List Bullet")
document.add_paragraph("first item in ordered list", style="List Number")

document.add_page_break()

# Table
document.add_heading("Table", level=1)

headers = ["id", "url", "name", "short_name", "birthday", "job"]
rows = [
    [
        "#1",
        "http://amiller.example.com",
        "Andrew Miller",
        "amiller",
        "11 December",
        "Testing Engineer",
    ],
    [
        "#2",
        "http://ataylor.example.com",
        "Anthony Taylor",
        "ataylor",
        "17 July",
        "Software Engineer",
    ],
    [
        "#3",
        "http://dmoore.example.com",
        "Daniel Moore",
        "dmoore",
        "2 March",
        "Testing Engineer",
    ],
    [
        "#4",
        "http://dsmith.example.com",
        "David Smith",
        "dsmith",
        "5 January",
        "Testing Engineer",
    ],
    [
        "#5",
        "http://awilson.example.com",
        "Alexander Wilson",
        "awilson",
        "11 April",
        "Software Engineer",
    ],
]

# Без style='Table Grid' у таблицы не будет выделена решетка
table = document.add_table(rows=1, cols=len(headers), style="Table Grid")
heading_cells = table.rows[0].cells

for i, value in enumerate(headers):
    # heading_cells[i].text = value.title()

    # Bold column
    heading_cells[i].paragraphs[0].add_run(value.title()).bold = True

for row in rows:
    row_cells = table.add_row().cells

    for i, value in enumerate(row):
        row_cells[i].text = value

document.add_page_break()

# Image
file_name_image = "image.jpg"

document.add_heading("Image", level=1)

document.add_heading("image__original", level=2)
document.add_picture(file_name_image)

document.add_heading("image__Inches_3", level=2)
document.add_picture(file_name_image, width=Inches(3.0))

document.add_heading("image__millimeters_30", level=2)
document.add_picture(file_name_image, width=Mm(30))

document.add_heading("image__centimeters_5", level=2)
document.add_picture(file_name_image, width=Cm(5))

document.add_page_break()

# Save
file_name_doc = "word.docx"
document.save(file_name_doc)

# Open file
os.startfile(file_name_doc)
