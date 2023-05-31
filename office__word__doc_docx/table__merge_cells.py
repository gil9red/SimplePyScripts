#!/usr/bin/env python3
# -*- coding: utf-8 -*-

__author__ = "ipetrash"


import os

# pip install python-docx
import docx


COL_COUNT = 5
ROW_COUNT = 5

HEADERS = ["TABLE HEADER", "", "", "", ""]
ROWS = [
    [f"{i}x{j}" for j in range(1, COL_COUNT + 1)]
    for i in range(1, ROW_COUNT + 1)
]


def fill_table(document, headers, rows):
    table = document.add_table(rows=1, cols=len(headers), style="Table Grid")
    heading_cells = table.rows[0].cells

    for i, value in enumerate(headers):
        # heading_cells[i].text = value

        # Bold column
        if value:
            heading_cells[i].paragraphs[0].add_run(value).bold = True

    for row in rows:
        cells = table.add_row().cells

        for i, value in enumerate(row):
            cells[i].text = value

    return table


document = docx.Document()
table_1 = fill_table(document, HEADERS, ROWS)

document.add_paragraph()

table_2 = fill_table(document, HEADERS, ROWS)

# Merge table header
table_2.cell(0, 0).merge(table_2.cell(0, COL_COUNT - 1))

# Merge rows 1x0 - 5x0
table_2.cell(1, 0).merge(table_2.cell(ROW_COUNT, 0))

# Merge cols 2x2 - 2x4
table_2.cell(2, 2).merge(table_2.cell(2, 4))

# Save
file_name_doc = "word.docx"
document.save(file_name_doc)

# Open file
os.startfile(file_name_doc)
