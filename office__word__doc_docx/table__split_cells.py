#!/usr/bin/env python3
# -*- coding: utf-8 -*-

__author__ = "ipetrash"


import os

# pip install python-docx
import docx


headers = ["VALUE"]
rows = [[""], [""]]


document = docx.Document()

table = document.add_table(rows=1, cols=len(headers), style="Table Grid")
heading_cells = table.rows[0].cells

for i, value in enumerate(headers):
    # heading_cells[i].text = value

    # Bold column
    heading_cells[i].paragraphs[0].add_run(value).bold = True

for row in rows:
    cells = table.add_row().cells

    for i, value in enumerate(row):
        cells[i].text = value

# Row 1
sub_table = table.cell(1, 0).add_table(rows=1, cols=3)
sub_table.style = "Table Grid"

for i, col in enumerate(sub_table.row_cells(0), 1):
    col.text = str(2**i)

# Row 2
sub_table = table.cell(2, 0).add_table(rows=1, cols=3)
for i, col in enumerate(sub_table.row_cells(0), 1):
    col.text = str(2**i)


# Save
file_name_doc = "word.docx"
document.save(file_name_doc)

# Open file
os.startfile(file_name_doc)
