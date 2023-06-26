#!/usr/bin/env python3
# -*- coding: utf-8 -*-

__author__ = "ipetrash"


"""
Меню находится в определенных письмах, приложенное как docx-файл.
Скрипт считывает последнее письмо, вытаскивает и сохраняет приложенный файл и
считывает из него таблицы меню и отображает таблицы на веб странице.

Устанавливать docx так: pip install python-docx

"""


import email
import imaplib
import os
import re
import sys
import traceback

from datetime import datetime, date

# pip install python-docx
from docx import Document
from flask import Flask, render_template_string, request

import config


if config.debug:
    import logging
    logging.basicConfig(level=logging.DEBUG, stream=sys.stdout)


def get_current_lunch_file_name():
    filename = date.today().strftime("%d.%m.%y") + ".docx"
    return os.path.join("lunch menu", filename)


def save_attachment(msg):
    """
    Given a message, save its attachments to the specified
    download folder (default is /tmp)

    return: file path to attachment
    """

    file_name = get_current_lunch_file_name()
    os.makedirs(os.path.dirname(file_name), exist_ok=True)

    for part in msg.walk():
        if part.get_content_maintype() == "multipart":
            continue

        if part.get("Content-Disposition") is None:
            continue

        if not os.path.exists(file_name):
            with open(file_name, "wb") as fp:
                fp.write(part.get_payload(decode=True))
        else:
            logging.debug("Lunch menu file exist: %s", file_name)

    return file_name


def add_lunch_email_info(msg, file_name):
    """
    Функция для добавления в поле comments docx информации о письме.

    :param msg:
    :param file_name:
    :return:
    """

    try:
        # Получаем заголовок письма
        data, charset = email.header.decode_header(msg["Subject"])[0]
        subject = data.decode(charset).strip().replace("\t", " ")
        logging.debug('Subject: "%s".', subject)

        # Получаем дату получения (отправления??) письма
        date_tuple = email.utils.parsedate_tz(msg["Date"])
        email_date = datetime.fromtimestamp(email.utils.mktime_tz(date_tuple))
        email_date = email_date.strftime(config.header_date_format)
        logging.debug("Date: %s.", email_date)

        # Открываем документ и переписываем поле comments, добавляя в него свою информацию
        document = Document(file_name)
        document.core_properties.comments = subject + "\t" + email_date
        document.save(file_name)

    except BaseException as e:
        logging.warning('add_lunch_email_info: "%s".', e)
        logging.warning(traceback.format_exc())


def get_last_lunch_menu():
    """
    Функция получает последнее письмо от указанного емейла
    и сохраняет из него приложенный файл.

    """

    # Проверка на то, что меню уже скачано
    file_name = get_current_lunch_file_name()
    if os.path.exists(file_name):
        logging.debug("Lunch menu file already exist: %s", file_name)
        return file_name

    logging.debug("Check last email.")

    connect = imaplib.IMAP4(config.smtp_server)
    connect.login(config.username, config.password)
    connect.select()

    logging.debug("Search emails from %s.", config.lunch_email)

    typ, msgnums = connect.search(None, f"(HEADER From {config.lunch_email})")
    last_id = msgnums[0].split()[-1]
    typ, data = connect.fetch(last_id, "(RFC822)")

    msg = email.message_from_bytes(data[0][1])

    logging.debug("Save lunch file name: %s.", file_name)
    file_name = save_attachment(msg)
    add_lunch_email_info(msg, file_name)

    connect.close()
    connect.logout()

    return file_name


app = Flask(__name__)

# Регулярка для поиска последовательностей пробелов: от двух подряд и более
multi_space_pattern = re.compile(r"[ ]{2,}")


def get_info_lunch_menu():
    rows = list()

    file_name = get_last_lunch_menu()
    if file_name is None:
        return rows

    logging.debug("Read lunch file name: %s.", file_name)

    document = Document(file_name)

    subject, date = "", ""
    comments = document.core_properties.comments
    if comments:
        parts = comments.split("\t", 2)
        if len(parts) == 2:
            subject, date = parts

    logging.debug('Subject: "%s".', subject)
    logging.debug('Date: "%s".', date)

    for table in document.tables:
        # Перебор начинается со второй строки, потому что, первая строка таблицы -- это строка "Обеденное меню"
        for row in table.rows[1:]:
            name, weight, price = [
                multi_space_pattern.sub(" ", i.text.strip()) for i in row.cells
            ]

            if name == weight == price or (not weight or not price):
                name = name.title()
                logging.debug(name)
                rows.append((name,))
                continue

            rows.append((name, weight, price))
            logging.debug(f"{name} {weight} {price}")

        # Таблицы в меню дублируются
        break

    return rows, subject, date


@app.route("/")
def index():
    logging.debug("/index from %s.", request.remote_addr)

    try:
        rows, subject, date = get_info_lunch_menu()
    except BaseException as e:
        logging.error(e)
        logging.error(traceback.format_exc())
        return f'Error: "{e}".', 500

    return render_template_string(
        """\
    <html>
    <head>
        <title>{{ subject }}</title>
        <link rel="stylesheet" href="static/style.css">
    </head>

    <body>

    <table>
        <tr><th colspan="3" align="center">{{ subject }} <small>(Меню получено: {{ date }})</small></th></tr>
        <tr><th>Название</th><th>Вес</th><th>Цена</th></tr>
        {% for row in rows %}

            {% if row|length == 3 %}
                <tr>
                    <td>{{ row[0] }}</td>
                    <td>{{ row[1] }}</td>
                    <td>{{ row[2] }}</td>
                </tr>
            {% else %}
                <tr>
                    <td class="category" colspan="3">{{ row[0] }}</td>
                </tr>
            {% endif %}
        {% endfor %}
    </table>

    </body>
    </html>
    """,
        rows=rows,
        subject=subject,
        date=date,
    )


if __name__ == "__main__":
    # Localhost
    app.run(port=5001)

    # # Public IP
    # app.run(host='0.0.0.0')
