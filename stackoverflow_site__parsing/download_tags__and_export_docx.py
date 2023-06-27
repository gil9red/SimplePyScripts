#!/usr/bin/env python3
# -*- coding: utf-8 -*-

__author__ = "ipetrash"


# ANALOG: https://github.com/gil9red/SimplePyScripts/blob/4d226491338b311fba75a841735f02c4f64c7379/stackoverflow/Py-StackExchange__examples/download_tags_excerpt__and_export_docx.py


import os
import re

from download_tags import get_all_tags

# pip install python-docx
import docx
from docx.shared import Cm


# SOURCE: https://github.com/gil9red/SimplePyScripts/blob/c4320a990755c78f5e4291a8448ac35cf7c2b250/office__word__doc_docx/hyperlink.py
def add_hyperlink(paragraph, url, text, color="1111FF", underline=True):
    """
    A function that places a hyperlink within a paragraph object.

    :param paragraph: The paragraph we are adding the hyperlink to.
    :param url: A string containing the required url
    :param text: The text displayed for the url
    :param color: color
    :param underline: underline
    :return: The hyperlink object
    """

    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(
        url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True
    )

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement("w:hyperlink")
    hyperlink.set(
        docx.oxml.shared.qn("r:id"),
        r_id,
    )

    # Create a w:r element
    new_run = docx.oxml.shared.OxmlElement("w:r")

    # Create a new w:rPr element
    rPr = docx.oxml.shared.OxmlElement("w:rPr")

    # Add color if it is given
    if color:
        c = docx.oxml.shared.OxmlElement("w:color")
        c.set(docx.oxml.shared.qn("w:val"), color)
        rPr.append(c)

    # Remove underlining if it is requested
    if not underline:
        u = docx.oxml.shared.OxmlElement("w:u")
        u.set(docx.oxml.shared.qn("w:val"), "none")
        rPr.append(u)

    u = docx.oxml.shared.OxmlElement("w:u")
    u.set(docx.oxml.shared.qn("w:val"), "single")
    rPr.append(u)

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)

    return hyperlink


# # Parse all pages:
# tags = get_all_tags()
tags = get_all_tags(need_pages=20)

headers = ["#", "NAME", "DESCRIPTION"]
rows = []

# Sorted tag
for i, name in enumerate(sorted(tags), 1):
    tag = tags[name]

    description = tag["description"]
    description = re.sub(" {2,}", " ", description)
    description = re.sub("\n{2,}", "\n", description)
    description = re.sub("\t{2,}", "\t", description)

    rows.append((i, name, tag["url_info"], description))

document = docx.Document()
table = document.add_table(rows=1, cols=len(headers), style="Table Grid")
heading_cells = table.rows[0].cells

for i, value in enumerate(headers):
    heading_cells[i].paragraphs[0].add_run(value).bold = True

for row in rows:
    cells = table.add_row().cells

    i, name, url, description = row

    cells[0].text = str(i)

    p = cells[1].paragraphs[0]
    add_hyperlink(p, url, name)

    cells[2].text = description

# Column size
column_size = (Cm(1.5), Cm(5.0), Cm(19.0))

for i, size in enumerate(column_size):
    for cell in table.columns[i].cells:
        cell.width = size

# Save
file_name_doc = "word.docx"
document.save(file_name_doc)

# Open file
os.startfile(file_name_doc)
